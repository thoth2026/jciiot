"""Read Document Skill — extract text + images from .docx, analyze with VLM."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from robot_agent.core.types import ExecutionContext, SkillResult
from robot_agent.skills.base import BaseSkill

logger = logging.getLogger(__name__)


def extract_docx_payload(
    path: str | Path,
    *,
    use_vision: bool = True,
    vision_base_url: str = "http://localhost:11434",
    vision_model: str = "qwen3-vl:8b",
    api_type: str = "ollama",
    api_key: str = "",
    vision_prompt: str | None = None,
) -> dict[str, Any]:
    """Extract prompt-ready text, tables, and optional image descriptions."""
    docx_path = Path(path)
    if not docx_path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    from docx import Document

    doc = Document(str(docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    tables: list[list[list[str]]] = []
    table_markdown: list[str] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
            table_markdown.append(_table_to_markdown(rows))

    images = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            name = rel.target_ref.split("/")[-1] if rel.target_ref else "image.png"
            images[name] = rel.target_part.blob

    img_descriptions = {}
    if use_vision and images:
        try:
            from robot_agent.core.vision_client import ask_vision

            prompt = vision_prompt or (
                "Describe this JCIIOT electronics factory SOP image. Identify "
                "stations, production lines, material bins/boxes, obstacles, "
                "navigation hints, and any pick/place relationship visible in "
                "the image. Be concise and factual."
            )
            for name, img_data in images.items():
                try:
                    img_descriptions[name] = ask_vision(
                        prompt,
                        img_data,
                        base_url=vision_base_url,
                        model=vision_model,
                        api_type=api_type,
                        api_key=api_key,
                    )
                except Exception as exc:
                    img_descriptions[name] = f"VLM error: {exc}"
        except Exception as exc:
            logger.warning("Vision analysis skipped: %s", exc)

    return {
        "file": str(docx_path),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "image_count": len(images),
        "images_analyzed": len(img_descriptions),
        "paragraphs": paragraphs,
        "text": full_text,
        "tables": tables,
        "table_markdown": "\n\n".join(table_markdown),
        "image_descriptions": img_descriptions,
    }


def _table_to_markdown(rows: list[list[str]]) -> str:
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


class ReadDocumentSkill(BaseSkill):
    """Read a .docx file: extract text, optionally describe images via VLM.

    LLM can invoke this with::

        {"skill_name": "read_document",
         "inputs": {"file": "knowledge/JCIIOT_2026_case_1_SOP.docx",
                    "use_vision": true}}
    """

    def __init__(
        self,
        *,
        ollama_base_url: str = "http://localhost:11434",
        vision_model: str = "qwen3-vl:8b",
        api_type: str = "ollama",
        api_key: str = "",
    ) -> None:
        super().__init__(
            name="read_document",
            description="Read .docx files, extract text and analyze images with vision model",
            keywords=("read", "document", "docx", "analyze", "vision", "parse", "extract"),
        )
        self._ollama_url = ollama_base_url
        self._vision_model = vision_model
        self._api_type = api_type
        self._api_key = api_key

    def run(self, context: ExecutionContext) -> SkillResult:
        file_path = context.metadata.get("inputs", {}).get("file", "")
        use_vision = context.metadata.get("inputs", {}).get("use_vision", True)

        path = Path(file_path)
        if not path.exists():
            return SkillResult(
                skill_name=self.name, success=False,
                message=f"File not found: {file_path}",
                payload={"file": file_path},
            )

        try:
            payload = extract_docx_payload(
                path,
                use_vision=use_vision,
                vision_base_url=self._ollama_url,
                vision_model=self._vision_model,
                api_type=self._api_type,
                api_key=self._api_key,
            )

            return SkillResult(
                skill_name=self.name,
                success=True,
                message=(
                    f"Read {payload['paragraph_count']} paragraphs, "
                    f"{payload['table_count']} tables, {payload['image_count']} images"
                    + (
                        f", {payload['images_analyzed']} analyzed by VLM"
                        if payload["image_descriptions"] else ""
                    )
                ),
                payload=payload,
            )
        except Exception as exc:
            logger.exception("ReadDocumentSkill failed")
            return SkillResult(
                skill_name=self.name, success=False,
                message=f"Failed: {exc}",
                payload={"file": file_path, "error": str(exc)},
            )
